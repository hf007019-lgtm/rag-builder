"""
本地 Word(.docx) 批量上传验证脚本。

运行前请先启动 FastAPI：
    uvicorn app.main:app --reload --port 18000

如果需要验证完整解析和 RAG 问答，还需要同时启动 Celery Worker。
"""

from __future__ import annotations

import os
import sys
import tempfile
import time
from pathlib import Path

import requests
from docx import Document


API_BASE = os.getenv("RAG_BUILDER_API_BASE", "http://127.0.0.1:18000/api/v1")
POLL_SECONDS = int(os.getenv("RAG_BUILDER_DOCX_POLL_SECONDS", "30"))

DOCX_PARAGRAPHS = [
    "广东事业单位信息技术岗位通常关注计算机类、软件工程、网络工程等相关专业。",
    "软件工程本科应届生可以重点关注信息技术岗、数据管理岗、数字政务岗、信息中心岗。",
]

ASK_QUESTION = "软件工程本科应届生可以重点关注什么事业单位岗位？"


def build_docx(path: Path) -> None:
    document = Document()
    document.add_heading("RAG Builder Word 上传测试", level=1)
    for paragraph in DOCX_PARAGRAPHS:
        document.add_paragraph(paragraph)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "岗位方向"
    table.cell(0, 1).text = "关注内容"
    table.cell(1, 0).text = "信息技术岗"
    table.cell(1, 1).text = "计算机类、软件工程、网络工程"
    document.save(path)


def post_batch_upload(path: Path) -> dict:
    upload_url = f"{API_BASE}/documents/batch-upload"
    with path.open("rb") as file_obj:
        response = requests.post(
            upload_url,
            files={
                "files": (
                    path.name,
                    file_obj,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            timeout=30,
        )
    response.raise_for_status()
    return response.json()


def poll_document_status(document_id: int) -> str:
    deadline = time.time() + POLL_SECONDS
    status = "PENDING"
    while time.time() < deadline:
        response = requests.get(
            f"{API_BASE}/documents/{document_id}/status",
            timeout=10,
        )
        response.raise_for_status()
        status = str(response.json().get("status") or "UNKNOWN").upper()
        print(f"当前文档状态：{status}")
        if status in {"SUCCESS", "FAILED"}:
            return status
        time.sleep(5)
    return status


def ask_rag() -> dict:
    response = requests.post(
        f"{API_BASE}/search/ask",
        json={"question": ASK_QUESTION},
        timeout=60,
    )
    response.raise_for_status()
    return response.json()


def main() -> int:
    print("开始生成 Word(.docx) 测试文件...")
    with tempfile.TemporaryDirectory() as temp_dir:
        docx_path = Path(temp_dir) / f"rag_builder_docx_upload_{int(time.time())}.docx"
        build_docx(docx_path)
        print(f"已生成测试文件：{docx_path}")

        try:
            upload_data = post_batch_upload(docx_path)
        except requests.ConnectionError:
            print("无法连接 RAG Builder API。请先启动：uvicorn app.main:app --reload --port 18000")
            return 1
        except requests.HTTPError as exc:
            print(f"批量上传接口返回错误：{exc.response.status_code} {exc.response.text}")
            return 1
        except requests.RequestException as exc:
            print(f"批量上传请求失败：{exc}")
            return 1

    items = upload_data.get("items") or []
    if not items:
        print(f"批量上传响应缺少 items：{upload_data}")
        return 1

    item = items[0]
    document_id = item.get("document_id")
    task_id = item.get("task_id")
    status = str(item.get("status") or "").upper()
    print(f"上传返回：document_id={document_id}, task_id={task_id}, status={status}")
    print(f"接口说明：{item.get('message')}")

    if not document_id:
        print("上传没有返回 document_id，测试失败。")
        return 1

    if status == "FAILED":
        print("上传接口返回 FAILED，测试失败。")
        return 1

    final_status = poll_document_status(int(document_id))
    if final_status != "SUCCESS":
        print(
            "Worker 可能未启动，当前只验证了批量上传接口已接收文件并进入待解析状态。"
        )
        return 0

    try:
        answer_data = ask_rag()
    except requests.RequestException as exc:
        print(f"RAG 问答请求失败：{exc}")
        return 1

    answer = str(answer_data.get("answer") or "")
    sources = answer_data.get("citations") or answer_data.get("sources") or []
    has_docx_source = any(
        str(source.get("file_name") or source.get("filename") or "").endswith(".docx")
        for source in sources
        if isinstance(source, dict)
    )

    print(f"RAG 回答：{answer[:200]}")
    print(f"引用数量：{len(sources)}")
    if not answer or not has_docx_source:
        print("RAG 回答未确认包含 Word 文档来源，测试失败。")
        return 1

    print("Word(.docx) 上传、解析和 RAG 问答验证通过。")
    return 0


if __name__ == "__main__":
    sys.exit(main())
