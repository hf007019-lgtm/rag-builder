# 从 io 导入 BytesIO
# 作用：把 MinIO 读取到的 bytes 包装成文件流，方便 pypdf 读取
from io import BytesIO


# 从 pathlib 导入 Path
# 作用：方便获取文件后缀名，例如 .txt、.pdf
from pathlib import Path


# 从 pypdf 导入 PdfReader
# 作用：读取 PDF 文件并提取文本
from pypdf import PdfReader


# 解析 txt 文本
# content：文件二进制内容
def parse_txt(content: bytes) -> str:

    # 优先使用 utf-8-sig 解码
    # 作用：兼容常见 UTF-8 文本，也能处理 BOM 头
    try:
        return content.decode("utf-8-sig")

    # 如果 utf-8 解码失败
    except UnicodeDecodeError:

        # 使用 gbk 解码
        # 作用：兼容部分 Windows 中文 txt 文件
        return content.decode("gbk", errors="ignore")


# 解析 PDF 文本
# content：PDF 文件二进制内容
def parse_pdf(content: bytes) -> str:

    # 把 bytes 包装成文件流
    # 作用：PdfReader 需要类似文件的对象
    pdf_stream = BytesIO(content)

    # 创建 PDF 读取器
    reader = PdfReader(pdf_stream)

    # 创建空列表
    # 作用：保存每一页提取出来的文本
    pages_text = []

    # 遍历 PDF 的每一页
    for page_index, page in enumerate(reader.pages):

        # 提取当前页文本
        page_text = page.extract_text() or ""

        # 如果当前页有文字
        if page_text.strip():

            # 加入页码标记
            # 作用：后面 sources 里能看出文本大概来自哪一页
            pages_text.append(f"【第 {page_index + 1} 页】\n{page_text}")

    # 把所有页面文本拼接起来
    full_text = "\n\n".join(pages_text)

    # 返回 PDF 全文
    return full_text


# 解析 Word .docx 文本
# content：Word 文件二进制内容
def parse_docx(content: bytes) -> str:

    # 延迟导入 python-docx
    # 作用：只有解析 Word 时才加载依赖
    from docx import Document as DocxDocument

    # 把 bytes 包装成文件流
    # 作用：python-docx 可以直接读取类文件对象
    docx_stream = BytesIO(content)

    # 创建 Word 文档对象
    document = DocxDocument(docx_stream)

    # 创建文本片段列表
    parts = []

    # 提取段落文本
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # 提取表格文本
    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    line.strip()
                    for line in cell.text.splitlines()
                    if line.strip()
                )
                if cell_text:
                    cells.append(cell_text)
            if cells:
                parts.append(" | ".join(cells))

    # 拼接全文
    full_text = "\n".join(parts)

    # 如果 Word 没有提取到有效文本，给出中文错误
    if not full_text.strip():
        raise ValueError("Word 文档未提取到有效文本。")

    # 返回 Word 全文
    return full_text


# 根据文件名和文件内容自动解析文本
# file_name：文件名
# content：文件二进制内容
def parse_document_content(file_name: str, content: bytes) -> str:

    # 获取文件后缀，并转成小写
    suffix = Path(file_name).suffix.lower()

    # 如果是 txt 或 markdown 文件
    if suffix in {".txt", ".md"}:

        # 调用文本解析函数
        return parse_txt(content)

    # 如果是 pdf 文件
    if suffix == ".pdf":

        # 调用 PDF 解析函数
        return parse_pdf(content)

    # 如果是 Word .docx 文件
    if suffix == ".docx":

        # 调用 Word 解析函数
        return parse_docx(content)

    # 如果是不支持的文件类型，直接报错
    raise ValueError(f"暂不支持的文件类型：{suffix}")
